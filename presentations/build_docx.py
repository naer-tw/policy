"""Convert the Markdown script to a properly-formatted Word .docx file."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

NAVY = RGBColor(0x0B, 0x2B, 0x4A)
GOLD = RGBColor(0xB8, 0x8E, 0x2F)
RED = RGBColor(0xB2, 0x3A, 0x3A)
GRAY = RGBColor(0x55, 0x55, 0x55)
INK = RGBColor(0x1A, 0x1A, 0x1A)
LIGHT_GRAY_BG = "F4F1E8"
NAVY_BG = "0B2B4A"
GOLD_BG = "F5EBC8"
RED_BG = "F8E6E6"

FONT = "微軟正黑體"
FONT_SCRIPTURE = "標楷體"


def set_run_font(run, font=FONT):
    """Force CJK font for a run."""
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)
    rFonts.set(qn("w:cs"), font)


def shade(element, fill_hex):
    """Add background shading to a paragraph or cell."""
    pPr = element if element.tag.endswith("}tcPr") else element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def add_run(p, text, size=11, bold=False, color=INK, font=FONT, italic=False):
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    set_run_font(r, font)
    return r


def add_p(doc, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4,
          line_spacing=1.4):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    return p


def add_heading(doc, text, level=1, color=NAVY, size=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = add_p(doc, align=align, space_before=18, space_after=10, line_spacing=1.3)
    if size is None:
        size = {1: 22, 2: 18, 3: 15, 4: 13}.get(level, 12)
    add_run(p, text, size=size, bold=True, color=color)
    return p


def add_separator(doc):
    p = add_p(doc, space_before=8, space_after=8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), "C9A247")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_callout(doc, lines, fill=LIGHT_GRAY_BG, accent=NAVY):
    """A shaded callout block with side accent."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    # background
    tcPr = cell._tc.get_or_add_tcPr()
    shade(tcPr, fill)
    # left border accent
    tcBorders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "36")
    left.set(qn("w:color"), "{:02X}{:02X}{:02X}".format(*accent))
    tcBorders.append(left)
    tcPr.append(tcBorders)

    # clear default paragraph
    cell.text = ""
    for i, (text, opts) in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_before = Pt(opts.get("space_before", 2))
        p.paragraph_format.space_after = Pt(opts.get("space_after", 4))
        p.paragraph_format.line_spacing = opts.get("line_spacing", 1.45)
        p.alignment = opts.get("align", WD_ALIGN_PARAGRAPH.LEFT)
        add_run(p, text,
                size=opts.get("size", 11),
                bold=opts.get("bold", False),
                color=opts.get("color", INK),
                font=opts.get("font", FONT),
                italic=opts.get("italic", False))


def add_speech(doc, paragraphs):
    """A speech block: pastor's word-for-word lines in a panel."""
    items = []
    for para, opts in paragraphs:
        items.append((para, opts))
    add_callout(doc, items, fill=GOLD_BG, accent=GOLD)


def add_meta_table(doc, rows):
    """Compact 2-column metadata table (label / value)."""
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Cm(3.5)
    tbl.columns[1].width = Cm(12.5)
    for i, (label, value) in enumerate(rows):
        cells = tbl.rows[i].cells
        cells[0].width = Cm(3.5)
        cells[1].width = Cm(12.5)
        shade(cells[0]._tc.get_or_add_tcPr(), "0B2B4A")
        for c, txt, color, bold in [
            (cells[0], label, RGBColor(0xFF, 0xFF, 0xFF), True),
            (cells[1], value, INK, False),
        ]:
            c.text = ""
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            add_run(p, txt, size=10, bold=bold, color=color)


# =============================================================
# Build document
# =============================================================
doc = Document()

# page setup
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)

# default style
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(11)
rpr = style.element.get_or_add_rPr()
rFonts = rpr.find(qn("w:rFonts"))
if rFonts is None:
    rFonts = OxmlElement("w:rFonts")
    rpr.append(rFonts)
rFonts.set(qn("w:eastAsia"), FONT)
rFonts.set(qn("w:ascii"), FONT)
rFonts.set(qn("w:hAnsi"), FONT)

# =============================================================
# COVER
# =============================================================
p = add_p(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=4)
add_run(p, "國教行動聯盟  Action Alliance on Basic Education", size=11, color=GRAY)

p = add_p(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=18, space_after=8)
add_run(p, "跨出牆外的服事", size=36, bold=True, color=NAVY)

p = add_p(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=6)
add_run(p, "從家庭崩解的呼喊", size=18, color=NAVY)
p = add_p(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=24)
add_run(p, "到公共政策的呼召", size=18, color=NAVY)

add_callout(doc, [
    ("「約瑟……他的枝條探出牆外。」",
     {"size": 16, "bold": True, "color": NAVY, "font": FONT_SCRIPTURE,
      "align": WD_ALIGN_PARAGRAPH.CENTER, "space_before": 8}),
    ("—— 創世記 49:22",
     {"size": 11, "color": GRAY,
      "align": WD_ALIGN_PARAGRAPH.CENTER, "space_after": 8}),
], fill=GOLD_BG, accent=GOLD)

p = add_p(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=30, space_after=4)
add_run(p, "講者完整逐字稿  ·  30 分鐘  ·  牧者聚集", size=11, color=GRAY)

# metadata
p = add_p(doc, space_before=20, space_after=0)
add_meta_table(doc, [
    ("講題", "跨出牆外的服事 — 從家庭崩解的呼喊到公共政策的呼召"),
    ("時長", "30 分鐘（含影片開場 2 分鐘、小組討論 ×2、Q&A、結語禱告）"),
    ("聽眾", "牧者聚集"),
    ("投影片", "共 17 張"),
    ("經文錨點", "創 49:22、尼 4:6、箴 31:8、斯 4:14"),
])

doc.add_page_break()

# =============================================================
# USAGE NOTES
# =============================================================
add_heading(doc, "使用說明", level=2)
notes = [
    "此稿為排練用完整版，正式講述時可依現場氣氛微調。",
    "每段標註「【投影片 X】」對應 PPT 位置。",
    "標註「（停一秒）」處請刻意留白，讓話語落下。",
    "標註「▶ 動作」處為現場引導動作。",
    "全程預估字數約 4,500 字，朗讀速度約 150 字/分。",
]
for n in notes:
    p = add_p(doc, space_before=2, space_after=4)
    add_run(p, "・ ", size=11, color=GOLD, bold=True)
    add_run(p, n, size=11)

add_separator(doc)

# =============================================================
# Time block helper
# =============================================================
def time_block(label, content_fn):
    """Render a timecode section header + content."""
    p = add_p(doc, space_before=18, space_after=8)
    add_run(p, "⏱  ", size=14, color=GOLD, bold=True)
    add_run(p, label, size=14, bold=True, color=NAVY)
    content_fn()


def slide_marker(text):
    p = add_p(doc, space_before=8, space_after=4)
    add_run(p, f"【{text}】", size=12, bold=True, color=RED)


def action_line(text):
    p = add_p(doc, space_before=4, space_after=4)
    add_run(p, "▶ ", size=11, color=GOLD, bold=True)
    add_run(p, text, size=11, italic=True, color=GRAY)


def speech_paragraphs(lines):
    """Render speech as gold callout, each line a paragraph."""
    items = []
    for ln in lines:
        if isinstance(ln, tuple):
            text, kind = ln
        else:
            text, kind = ln, "normal"
        opts = {"size": 12, "line_spacing": 1.6,
                "space_before": 2, "space_after": 4}
        if kind == "pause":
            opts.update({"size": 10, "color": GRAY, "italic": True,
                         "align": WD_ALIGN_PARAGRAPH.CENTER})
        elif kind == "emphasis":
            opts.update({"bold": True, "color": NAVY, "size": 13})
        elif kind == "scripture":
            opts.update({"font": FONT_SCRIPTURE, "size": 14, "bold": True,
                         "color": NAVY, "align": WD_ALIGN_PARAGRAPH.CENTER})
        elif kind == "scripture_ref":
            opts.update({"size": 10, "color": GRAY,
                         "align": WD_ALIGN_PARAGRAPH.CENTER})
        items.append((text, opts))
    add_callout(doc, items, fill=GOLD_BG, accent=GOLD)


# =============================================================
# 1. 影片開場
# =============================================================
def block1():
    slide_marker("投影片 1：封面")
    action_line("簡短自我介紹（1 句話以內）後直接說：")
    speech_paragraphs([
        "「今天的時間很有限，我想先邀請各位牧者，跟我一起看一段兩分鐘的畫面。",
        "這是過去三年，國教行動聯盟在台灣這塊土地上做的事。請看。」",
    ])
    action_line("播放國教盟 2023-2025 倡議成果影片（2 分鐘）。")
time_block("00:00 — 02:00　影片開場（2 分鐘）", block1)

# =============================================================
# 2. 現場真實
# =============================================================
def block2a():
    add_heading(doc, "02:00 — 02:15　影片結束銜接", level=3, color=GOLD, size=12)
    action_line("影片結束、燈光亮起、停留 2 秒，再開口。")
    speech_paragraphs([
        "「各位牧者，您剛剛看見的，是過去三年神在『牆外』動工的痕跡。",
        "但今天，我要先帶各位回到兩個現場——",
        ("一個叫剴剴，", "emphasis"),
        ("一個叫性平法。", "emphasis"),
        "因為這兩個現場，正在告訴台灣教會：",
        ("我們再也不能只在牆內服事了。", "emphasis"),
    ])

    add_heading(doc, "02:15 — 03:45　剴剴（90 秒）", level=3, color=GOLD, size=12)
    slide_marker("投影片 2：剴剴")
    speech_paragraphs([
        "「2023 年的冬天，台北一個 1 歲多的男孩。為了保護家屬，我們叫他剴剴。",
        "他原本應該在保母家被照顧。社工有去訪視過——那本來是社會安全網最後一道防線。",
        "但剴剴沒有等到下一次訪視。",
        "他在那雙本該擁抱他的手裡，永遠閉上了眼睛。",
        ("（停一秒）", "pause"),
        ("弟兄姊妹，整個台灣震驚的，不只是一個孩子的死，而是震驚我們以為的『系統』，原來這麼薄。", "normal"),
        ("社工被起訴、一審判刑——這在台灣是史無前例的。", "normal"),
        "但判完之後呢？",
        ("（停一秒）", "pause"),
        ("下一個剴剴，會不會還在發生？」", "emphasis"),
    ])

    add_heading(doc, "03:45 — 05:15　性平法（90 秒）", level=3, color=GOLD, size=12)
    slide_marker("投影片 3：性平法")
    speech_paragraphs([
        "「另一個現場是性平法。",
        "各位牧者，這不只是某一個爭議議題，而是整個世代的價值權柄之爭——",
        ("誰來定義我們的孩子，應該長成什麼樣子？", "emphasis"),
        "過去兩年，我們看見家長在學校門口拿到的，是他們從來沒見過的教材。",
        "想去公聽會發言，被排除在程序之外。",
        "家長想問：『我的孩子在學什麼？』制度卻說：『這不是你的事。』",
        ("（停一秒）", "pause"),
        ("今年國教盟召開了『性平法再校準』公聽會，三大主張就是要把教育拉回教育的本位——讓家長重新有座位、讓孩子重新被守護。」", "normal"),
    ])

    add_heading(doc, "05:15 — 06:00　兩道呼喊（45 秒）", level=3, color=GOLD, size=12)
    slide_marker("投影片 4：這不是兩個個案，是兩道呼喊")
    speech_paragraphs([
        "「弟兄姊妹，這不是兩個個案，這是兩道呼喊。",
        "一個是『保護的牆破了』——",
        "一個是『價值的牆斜了』。",
        "這就是 2026 年台灣兒少與家庭的真實處境。",
        "而教會——我們是這個社會剩下最完整的一張網。",
        "但我們需要走出去。",
        ("（停一秒）", "pause"),
        ("教會還能只在牆內服事嗎？」", "emphasis"),
    ])
time_block("02:00 — 06:00　現場真實（4 分鐘）", block2a)

# =============================================================
# 3. 成果聚焦
# =============================================================
def block3():
    add_heading(doc, "06:00 — 07:00　三個數字（60 秒）", level=3, color=GOLD, size=12)
    slide_marker("投影片 5：三個數字")
    speech_paragraphs([
        "「但弟兄姊妹，我來不是要讓各位陷在絕望裡。我來是要告訴各位——神已經動工了。",
        "過去三年，神讓國教盟做的事，用三個數字濃縮：",
        ("113 篇新聞稿、286 場活動、168 個合作組織。", "emphasis"),
        "媒體聲量超過 7,917 則。",
        ("校園安全 8 項訴求，6 項政府已經回應。", "normal"),
        ("（停一秒）", "pause"),
        "這不是國教盟的功勞——",
        ("這是當有人願意跨出牆外，神就讓果子結出來。", "emphasis"),
    ])

    add_heading(doc, "07:00 — 08:00　尼希米的牆（60 秒）", level=3, color=GOLD, size=12)
    slide_marker("投影片 6：尼希米記 4:6")
    speech_paragraphs([
        "「尼希米記第四章第六節說：",
        ("『這樣，我們修造城牆，城牆就都連絡，因為百姓齊心做工。』", "scripture"),
        ("（停一秒）", "pause"),
        "尼希米重建城牆的時候，沒有一個人能單獨修完整道牆。",
        "每一段都靠不同的家族、不同的職業、彼此緊鄰著做。",
        "國教盟這三年做的，就是這件事——",
        ("把『願意修牆的人』連起來。", "emphasis"),
        ("而今天，我也要邀請台灣的教會，加入這道城牆的修造。」", "emphasis"),
    ])
time_block("06:00 — 08:00　成果聚焦（2 分鐘）", block3)

# =============================================================
# 4. 小組討論 1
# =============================================================
def block4():
    slide_marker("投影片 7：討論題 1")
    action_line("投影出題後，主持人說：")
    speech_paragraphs([
        "「在我們進入下半場之前，我想先請各位牧者，停下來一分鐘。",
        "不需要分析、不需要對策，只要分享：",
        ("剛剛影片裡的孩子、剴剴的案件、性平法的張力——哪一個畫面最觸動你？", "emphasis"),
        "身為牧者，神正在如何透過這些『牆外』的呼喊，對你和你的會眾說話？",
        ("計時 3 分鐘，每位 30 秒就夠。請現在開始。」", "normal"),
    ])
    action_line("計時器 03:00 開始倒數。")
    action_line("計時到剩 30 秒時：「最後 30 秒，請最後一位牧者分享。」")
    action_line("時間到：「謝謝各位。我們繼續。」")
time_block("08:00 — 11:00　🟠 小組討論 1（3 分鐘）", block4)

# =============================================================
# 5. 2026 重點一
# =============================================================
def block5():
    add_heading(doc, "11:00 — 12:15　兒童家庭部（75 秒）", level=3, color=GOLD, size=12)
    slide_marker("投影片 8：兒童家庭部")
    speech_paragraphs([
        "「剛剛我們已經把感動接住了。",
        "接下來，我必須帶各位看見：這個時代神要教會切入的戰場在哪裡。",
        ("第一個戰場：兒童家庭部。", "emphasis"),
        "各位牧者，請看這三個數字——",
        ("兒少通報案件，五年內增加 65%。", "normal"),
        ("但實際處置率，只有 7%。", "normal"),
        "那台灣現在有幾個專責的兒少部會？",
        ("（停一秒）", "pause"),
        ("零個。", "emphasis"),
        ("（停一秒）", "pause"),
        "日本去年成立了『こども家庭庁』，韓國有保健福祉部專責處。",
        ("台灣呢？我們連一個專責部會都沒有。", "emphasis"),
        "國教盟正在推動的，是讓兒少權法第 7 條那個只是『會報』的層級，升格為二級『兒童家庭部』——",
        ("這不只是行政組織的調整，這是國家對下一代的態度宣告。」", "emphasis"),
    ])

    add_heading(doc, "12:15 — 13:30　兒少權法六大缺口（75 秒）", level=3, color=GOLD, size=12)
    slide_marker("投影片 9：六大缺口")
    speech_paragraphs([
        "「第二件事：兒少權法修法。",
        "衛福部預告的草案有 165 條，看起來很完整。",
        ("但國教盟逐條比對之後，找出六大缺口——", "normal"),
        "通報認定、社工保證人地位、安置責任、兒少代表權、資訊揭露、跨部會權責。",
        ("最關鍵的就是這個金色框框：社工保證人地位。", "emphasis"),
        ("（停一秒）", "pause"),
        ("剴剴案的社工被判刑之後，整個社工體系在崩潰。", "normal"),
        "為什麼？因為法律對她們的責任邊界沒有寫清楚。",
        "我們不是要保護任何一方。",
        ("我們是要把制度的牆——重新砌實。」", "emphasis"),
    ])

    add_heading(doc, "13:30 — 14:00　為不能自辯者（30 秒）", level=3, color=GOLD, size=12)
    slide_marker("投影片 10：箴言 31:8")
    speech_paragraphs([
        "「箴言 31 章 8 節：",
        ("『你當為不能自辯的開口，為一切孤獨的伸冤。』", "scripture"),
        ("（停一秒）", "pause"),
        "弟兄姊妹，兒少在公共政策的桌上，永遠是『不能自辯』的那一群。",
        ("今天教會若不為他們開口，誰開口？」", "emphasis"),
    ])
    p = add_p(doc, space_before=6, space_after=6)
    add_run(p, "14:00 — 15:00　", size=10, color=GRAY, bold=True)
    add_run(p, "（緩衝時間 — 視現場節奏，若前段超時可吸收 0-60 秒）", size=10, italic=True, color=GRAY)
time_block("11:00 — 15:00　2026 重點一（4 分鐘）", block5)

# =============================================================
# 6. 2026 重點二
# =============================================================
def block6():
    add_heading(doc, "15:00 — 16:30　家庭韌性聯盟（90 秒）", level=3, color=GOLD, size=12)
    slide_marker("投影片 11：家庭韌性行動聯盟")
    speech_paragraphs([
        "「除了兒少議題本身，2026 年國教盟還在兩個方向上跨出新的一步。",
        ("第一個方向是橫向——家庭韌性行動聯盟。", "emphasis"),
        "過去倡議的盲點是『議題分散』——婚家、生育、教養、托育、長照、信仰，各自為政。",
        "但家庭從來不是被『某一個政策』擊垮的。",
        ("家庭是被整個社會的擠壓——擠垮的。", "emphasis"),
        ("（停一秒）", "pause"),
        "所以我們要把這些議題縫合起來，從『單一議題』走向『家庭總體戰略』。",
        "而這件事，教會不能缺席。",
        ("因為家庭本來就是神所設立的秩序，沒有人比教會更知道家庭是什麼。」", "emphasis"),
    ])

    add_heading(doc, "16:30 — 18:00　藥物濫用國際研討會（90 秒）", level=3, color=GOLD, size=12)
    slide_marker("投影片 12：藥物濫用防制國際研討會")
    speech_paragraphs([
        "「第二個方向是縱向——藥物濫用防制國際研討會。",
        ("各位牧者，我們的孩子現在面對的，不是 30 年前的毒品。", "emphasis"),
        "依托咪酯、喪屍煙彈、新型尼古丁袋——",
        ("這些都是化學戰等級的攻擊。", "emphasis"),
        ("（停一秒）", "pause"),
        "今年國教盟參與藥物濫用防制國際研討會，把『唾液快篩二層檢測機制』帶上國際對話桌。",
        "我們不只是台灣內部的倡議組織，我們在國際對話桌上為台灣的孩子發聲。",
        "弟兄姊妹，這不是技術問題——",
        ("這是為這個世代築牆。」", "emphasis"),
    ])
time_block("15:00 — 18:00　2026 重點二（3 分鐘）", block6)

# =============================================================
# 7. 呼籲教會
# =============================================================
def block7():
    add_heading(doc, "18:00 — 19:30　公共政策小組（90 秒）", level=3, color=GOLD, size=12)
    slide_marker("投影片 13：四象限・公共政策小組")
    speech_paragraphs([
        "「講到這裡，我想正式向各位牧者發出邀請。",
        ("我邀請每一間教會，開始思考成立『公共政策小組』。", "emphasis"),
        "我知道每一位牧者都很忙——所以我不是邀請各位另起爐灶。",
        ("我是邀請各位，在你的教會體質裡，找一塊你最能委身的。", "emphasis"),
        ("（停一秒）", "pause"),
        "有些教會有禱告勇士——那就成為『代禱』那一塊：為立法者、社工、教師守望。",
        "有些教會有律師、學者、研究者——那就成為『研究』那一塊：聖經倫理 × 政策分析。",
        "有些教會有發言力、媒體連結——那就成為『發聲』那一塊：公聽會、媒體、聯署。",
        "有些教會有社工、輔導、家庭事工——那就成為『陪伴』那一塊：受害家庭、第一線工作者。",
        ("（停一秒，加重語氣）", "pause"),
        ("沒有任何一間教會需要做全部。", "emphasis"),
        ("但每一間教會，都可以做一塊。」", "emphasis"),
    ])

    add_heading(doc, "19:30 — 20:15　以斯帖（45 秒）", level=3, color=GOLD, size=12)
    slide_marker("投影片 14：以斯帖記 4:14")
    speech_paragraphs([
        "「以斯帖記 4 章 14 節：",
        ("『焉知你得了王后的位分，不是為現今的機會嗎？』", "scripture"),
        ("（停一秒）", "pause"),
        "以斯帖如果留在王宮裡，安全地當她的王后，整個猶太民族就沒了。",
        ("她的『位分』，是為了『跨出去』。", "emphasis"),
        "弟兄姊妹，台灣教會被神放在這個時代、這片土地——",
        ("不是為了我們的安全，", "emphasis"),
        ("是為了現今的機會。」", "emphasis"),
    ])

    add_heading(doc, "20:15 — 21:00　約瑟金句（45 秒｜全場高點）",
                level=3, color=RED, size=12)
    slide_marker("投影片 15：創世記 49:22")
    action_line("放慢語速，每一句之間至少停半秒。")
    speech_paragraphs([
        "「弟兄姊妹，創世記四十九章二十二節說：",
        ("『約瑟……他的枝條探出牆外。』", "scripture"),
        ("（停一秒）", "pause"),
        ("教會的根，留在泉旁——我們不離開基督。", "emphasis"),
        ("但教會的果子，必須探出牆外。", "emphasis"),
        ("（停一秒）", "pause"),
        "這就是國教盟的呼召，",
        ("也是今天我們對每一間教會的邀請。」", "emphasis"),
    ])

    p = add_p(doc, space_before=6, space_after=6)
    add_run(p, "21:00 — 22:00　", size=10, color=GRAY, bold=True)
    add_run(p, "（緩衝時間，視前段節奏，若略長則直接進入討論題 2）", size=10, italic=True, color=GRAY)
time_block("18:00 — 22:00　呼籲教會（4 分鐘）", block7)

# =============================================================
# 8. 小組討論 2
# =============================================================
def block8():
    slide_marker("投影片 16：討論題 2")
    action_line("投影出題後，主持人說：")
    speech_paragraphs([
        "「最後，我想用一個更具體的問題，邀請各位牧者把感動變成行動。",
        "如果您的教會要踏出第一步參與公共政策——",
        ("第一，是先從『代禱、研究、發聲、陪伴』哪一塊開始？", "emphasis"),
        ("第二，最大的攔阻是什麼？", "emphasis"),
        ("第三，您願意帶誰一起？", "emphasis"),
        ("（停一秒）", "pause"),
        "這次不需要全部回答，每位牧者挑一題回應就好。",
        "重點不是答得完整——",
        ("重點是讓你想到的那個人，等一下你會打給他。", "emphasis"),
        ("計時 3 分鐘，請現在開始。」", "normal"),
    ])
    action_line("計時器 03:00 開始倒數。")
    action_line("時間到：「謝謝各位。我們進入最後一段，回到大家分享。」")
time_block("22:00 — 25:00　🟠 小組討論 2（3 分鐘）", block8)

# =============================================================
# 9. 回饋分享 + Q&A
# =============================================================
def block9():
    slide_marker("投影片 16（維持）或可換至空白")
    action_line("邀請 2-3 位牧者分享：")
    speech_paragraphs([
        "「請 2-3 位牧者，跟我們分享：剛剛你們組裡，神感動最深的一個答案是什麼？」",
    ])

    add_heading(doc, "🎤 主持人「對焦提問」四句（每位分享後擇一回應，不評論、只對焦）",
                level=3, color=NAVY, size=12)

    focus_qs = [
        ("①  聚焦感動型",
         "「牧者，您剛剛說的那個畫面——神給您這個感動，您覺得是要先停下來禱告，還是已經到了該行動的時候？」",
         "對方分享感動但未定方向"),
        ("②  連結同行型",
         "「您提到那位（同工/長執/弟兄姊妹）——如果今天會後您打給他一通電話，您會怎麼開口？」",
         "對方提到具體人名"),
        ("③  教會體質型",
         "「這個負擔聽起來很像神放在您教會體質裡的呼召——您覺得這是您一個人在扛，還是神也在預備您的會眾？」",
         "對方分享教會方向"),
        ("④  第一小步型",
         "「從現在到下個主日，您願意先做的『最小那一步』是什麼？」",
         "對方分享意向但較抽象"),
    ]
    tbl = doc.add_table(rows=len(focus_qs), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    for i, (title, q, when) in enumerate(focus_qs):
        cells = tbl.rows[i].cells
        cells[0].width = Cm(4.0)
        cells[1].width = Cm(12.0)
        shade(cells[0]._tc.get_or_add_tcPr(), "F5EBC8")
        cells[0].text = ""
        cells[1].text = ""
        p = cells[0].paragraphs[0]
        p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(2)
        add_run(p, title, size=11, bold=True, color=NAVY)
        p2 = cells[0].add_paragraph()
        p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(4)
        add_run(p2, f"適用：{when}", size=9, italic=True, color=GRAY)

        p3 = cells[1].paragraphs[0]
        p3.paragraph_format.space_before = Pt(4); p3.paragraph_format.space_after = Pt(4)
        p3.paragraph_format.line_spacing = 1.5
        add_run(p3, q, size=11)

    p = add_p(doc, space_before=8, space_after=4)
    add_run(p, "▶ 使用原則：", size=11, color=GOLD, bold=True)
    add_run(p, "每位分享者只回應一句，不展開。讓會場保持禱告與默想的氛圍。", size=11)

    action_line("若還有時間，開放 1-2 題 Q&A：「我們還有大約一分鐘，誰有想要當場問的問題？」")
    action_line("若無人提問：「沒有問題也很好——有時候沉默就是聖靈在做工。我們直接進入禱告。」")
time_block("25:00 — 29:00　回饋分享 ＋ Q&A（4 分鐘）", block9)

# =============================================================
# 10. 結語禱告
# =============================================================
def block10():
    slide_marker("投影片 17：結語禱告")
    action_line("「請各位牧者，起立——或是閉上眼睛，我們一同來禱告。」")
    add_heading(doc, "禱告（約 50 秒，請慢、請穩）", level=3, color=GOLD, size=12)
    add_callout(doc, [
        ("主啊，",
         {"size": 13, "bold": True, "color": NAVY, "font": FONT_SCRIPTURE,
          "align": WD_ALIGN_PARAGRAPH.CENTER, "space_before": 6, "space_after": 8,
          "line_spacing": 1.6}),
        ("謝謝祢讓我們今天聽見了牆外的呼喊。",
         {"size": 12, "font": FONT_SCRIPTURE, "color": NAVY,
          "align": WD_ALIGN_PARAGRAPH.CENTER, "space_after": 6, "line_spacing": 1.6}),
        ("謝謝祢讓我們看見：剴剴的離開不是白白的，性平法的張力不是絕望的，",
         {"size": 12, "font": FONT_SCRIPTURE, "color": NAVY,
          "align": WD_ALIGN_PARAGRAPH.CENTER, "space_after": 0, "line_spacing": 1.6}),
        ("因為祢正在感動祢的教會起來。",
         {"size": 12, "font": FONT_SCRIPTURE, "color": NAVY,
          "align": WD_ALIGN_PARAGRAPH.CENTER, "space_after": 10, "line_spacing": 1.6}),
        ("主啊，求祢給台灣每一位牧者，今天回去之後——",
         {"size": 12, "font": FONT_SCRIPTURE, "color": NAVY,
          "align": WD_ALIGN_PARAGRAPH.CENTER, "space_after": 4, "line_spacing": 1.6}),
        ("至少打一通電話、找到一位同行的人。",
         {"size": 13, "bold": True, "font": FONT_SCRIPTURE, "color": NAVY,
          "align": WD_ALIGN_PARAGRAPH.CENTER, "space_after": 10, "line_spacing": 1.6}),
        ("求祢讓祢的教會，像約瑟那棵枝子——",
         {"size": 12, "font": FONT_SCRIPTURE, "color": NAVY,
          "align": WD_ALIGN_PARAGRAPH.CENTER, "space_after": 4, "line_spacing": 1.6}),
        ("根扎在祢裡面，",
         {"size": 13, "bold": True, "font": FONT_SCRIPTURE, "color": NAVY,
          "align": WD_ALIGN_PARAGRAPH.CENTER, "space_after": 0, "line_spacing": 1.6}),
        ("果子結在這片土地上。",
         {"size": 13, "bold": True, "font": FONT_SCRIPTURE, "color": NAVY,
          "align": WD_ALIGN_PARAGRAPH.CENTER, "space_after": 10, "line_spacing": 1.6}),
        ("讓我們的下一代，因為今天這群牧者的『跨出去』，得以看見祢的榮耀。",
         {"size": 12, "font": FONT_SCRIPTURE, "color": NAVY,
          "align": WD_ALIGN_PARAGRAPH.CENTER, "space_after": 12, "line_spacing": 1.6}),
        ("奉主耶穌的名禱告，",
         {"size": 12, "font": FONT_SCRIPTURE, "color": NAVY,
          "align": WD_ALIGN_PARAGRAPH.CENTER, "space_after": 4, "line_spacing": 1.6}),
        ("阿們。",
         {"size": 14, "bold": True, "font": FONT_SCRIPTURE, "color": NAVY,
          "align": WD_ALIGN_PARAGRAPH.CENTER, "space_after": 6, "line_spacing": 1.6}),
    ], fill=GOLD_BG, accent=GOLD)
    action_line("抬頭，微笑：「謝謝各位牧者。願主大大祝福您和您的教會。」")
time_block("29:00 — 30:00　結語禱告（1 分鐘）", block10)

doc.add_page_break()

# =============================================================
# 排練檢核表
# =============================================================
add_heading(doc, "📋  排練檢核表", level=1, color=NAVY, size=20)

add_heading(doc, "講者準備", level=3, color=NAVY)
for item in [
    "影片提前 30 分鐘測試播放（音量、畫面比例）",
    "計時器（手機/番茄鐘）放在講桌可見處",
    "確認簡報切換器（clicker）電池",
    "準備一杯溫水（不要冰）",
]:
    p = add_p(doc, space_before=1, space_after=2)
    add_run(p, "☐  ", size=11, color=GOLD, bold=True)
    add_run(p, item, size=11)

add_heading(doc, "節奏掌握", level=3, color=NAVY)
for item in [
    "影片結束後停 2 秒再開口（讓影片餘韻沉澱）",
    "金句經文（特別是 Slide 15）必須放慢",
    "兩次小組討論準時收尾，不要讓討論吃掉後面段落",
    "回饋分享時只用對焦句回應，避免評論延展",
]:
    p = add_p(doc, space_before=1, space_after=2)
    add_run(p, "☐  ", size=11, color=GOLD, bold=True)
    add_run(p, item, size=11)

add_heading(doc, "心境準備", level=3, color=NAVY)
for item in [
    "開場前默念：「我不是來傳遞資訊，我是來傳遞呼召。」",
    "兩個案例（剴剴、性平）要帶感情、不要帶情緒",
    "結語禱告前讓自己安靜 5 秒，讓聖靈帶領禱告詞",
]:
    p = add_p(doc, space_before=1, space_after=2)
    add_run(p, "☐  ", size=11, color=GOLD, bold=True)
    add_run(p, item, size=11)

# =============================================================
# 應變備案
# =============================================================
add_heading(doc, "🛟  應變備案", level=1, color=NAVY, size=20)

contingencies = [
    ("影片無法播放", "跳過影片，直接從「現場真實」開始，加長剴剴/性平描述"),
    ("討論過熱、超時", "主動截斷：「謝謝，請各組保留這個感動，我們繼續往下」"),
    ("牧者提出尖銳異議", "不辯論：「謝謝您的提醒，這是我們要持續對話的——會後我們再深談」"),
    ("時間嚴重落後", "砍掉 Slide 14（以斯帖）一段，直接從 Slide 13 跳 Slide 15"),
    ("時間提前結束", "多開放 1-2 題 Q&A，或邀請第 4 位牧者分享"),
]
tbl = doc.add_table(rows=len(contingencies)+1, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
# header
header = tbl.rows[0].cells
for i, h in enumerate(["狀況", "處理方式"]):
    shade(header[i]._tc.get_or_add_tcPr(), "0B2B4A")
    header[i].text = ""
    p = header[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, h, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
header[0].width = Cm(4.5)
header[1].width = Cm(11.5)
for i, (k, v) in enumerate(contingencies, 1):
    cells = tbl.rows[i].cells
    cells[0].width = Cm(4.5)
    cells[1].width = Cm(11.5)
    cells[0].text = ""; cells[1].text = ""
    add_run(cells[0].paragraphs[0], k, size=11, bold=True, color=NAVY)
    add_run(cells[1].paragraphs[0], v, size=11)
    for c in cells:
        c.paragraphs[0].paragraph_format.space_before = Pt(3)
        c.paragraphs[0].paragraph_format.space_after = Pt(3)

# =============================================================
# 時間配置總表
# =============================================================
add_heading(doc, "📐  時間配置總表", level=1, color=NAVY, size=20)

schedule = [
    ("1", "影片開場", "00:00", "02:00", "2"),
    ("2", "剴剴 ＋ 性平 ＋ 兩道呼喊", "02:00", "06:00", "4"),
    ("3", "三個數字 ＋ 尼希米", "06:00", "08:00", "2"),
    ("4", "🟠 小組討論 1", "08:00", "11:00", "3"),
    ("5", "兒童家庭部 ＋ 六大缺口 ＋ 箴言", "11:00", "15:00", "4"),
    ("6", "家庭韌性 ＋ 國際研討會", "15:00", "18:00", "3"),
    ("7", "公共政策小組 ＋ 以斯帖 ＋ 約瑟金句", "18:00", "22:00", "4"),
    ("8", "🟠 小組討論 2", "22:00", "25:00", "3"),
    ("9", "回饋分享 ＋ Q&A", "25:00", "29:00", "4"),
    ("10", "結語禱告", "29:00", "30:00", "1"),
]
tbl = doc.add_table(rows=len(schedule)+2, cols=5)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["#", "段落", "起", "訖", "分鐘"]
widths = [Cm(1.0), Cm(8.5), Cm(2.0), Cm(2.0), Cm(2.0)]
for i, h in enumerate(headers):
    cell = tbl.rows[0].cells[i]
    cell.width = widths[i]
    shade(cell._tc.get_or_add_tcPr(), "0B2B4A")
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, h, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
for r, row in enumerate(schedule, 1):
    cells = tbl.rows[r].cells
    for i, val in enumerate(row):
        cells[i].width = widths[i]
        cells[i].text = ""
        p = cells[i].paragraphs[0]
        align = WD_ALIGN_PARAGRAPH.CENTER if i != 1 else WD_ALIGN_PARAGRAPH.LEFT
        p.alignment = align
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        add_run(p, val, size=11, bold=(i == 0))
# total row
total_cells = tbl.rows[-1].cells
for i, val in enumerate(["", "總計", "", "", "30"]):
    total_cells[i].width = widths[i]
    shade(total_cells[i]._tc.get_or_add_tcPr(), "F5EBC8")
    total_cells[i].text = ""
    p = total_cells[i].paragraphs[0]
    align = WD_ALIGN_PARAGRAPH.CENTER if i != 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.alignment = align
    add_run(p, val, size=11, bold=True, color=NAVY)

# tail
p = add_p(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=30, space_after=6)
add_run(p, "願主祝福這次聚集，使祢的話語不徒然返回。",
        size=12, italic=True, color=NAVY, font=FONT_SCRIPTURE)
p = add_p(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=0)
add_run(p, "—— 國教行動聯盟  Action Alliance on Basic Education",
        size=10, color=GRAY)

out = "/home/user/policy/presentations/跨出牆外的服事_講者逐字稿.docx"
doc.save(out)
print(f"Saved: {out}")
